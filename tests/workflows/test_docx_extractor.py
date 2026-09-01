import io
import unittest
import zipfile
import sys
from pathlib import Path
import docx

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT / "services" / "docx-extractor"))

from app.validation import validate_docx_preflight, DocxValidationError
from app.extractor import extract_text_from_docx


def create_mock_docx_with_builder(builder_fn) -> bytes:
    doc = docx.Document()
    builder_fn(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def create_corrupt_or_custom_zip(extra_files: dict = None, omit_main: bool = False) -> bytes:
    doc = docx.Document()
    doc.add_paragraph("Base paragraph")
    buf = io.BytesIO()
    doc.save(buf)
    
    if extra_files or omit_main:
        src_zip = zipfile.ZipFile(buf, "r")
        new_buf = io.BytesIO()
        with zipfile.ZipFile(new_buf, "w", zipfile.ZIP_DEFLATED) as dst_zip:
            for item in src_zip.infolist():
                if omit_main and item.filename == "word/document.xml":
                    continue
                dst_zip.writestr(item, src_zip.read(item.filename))
            if extra_files:
                for fname, fcontent in extra_files.items():
                    dst_zip.writestr(fname, fcontent)
        return new_buf.getvalue()
    return buf.getvalue()


class TestDocxExtractor(unittest.TestCase):

    def test_docx_001_paragraphs(self):
        def build(doc):
            doc.add_paragraph("Título del Documento")
            doc.add_paragraph("Párrafo A de prueba")
            doc.add_paragraph("Párrafo B de prueba")
            
        raw = create_mock_docx_with_builder(build)
        res = extract_text_from_docx(raw)
        self.assertEqual(res["status"], "ok")
        self.assertIn("Título del Documento", res["text"])
        self.assertIn("Párrafo A de prueba", res["text"])
        self.assertIn("Párrafo B de prueba", res["text"])
        self.assertEqual(res["metadata"]["paragraph_count"], 3)
        self.assertEqual(res["metadata"]["table_count"], 0)

    def test_docx_002_tables(self):
        def build(doc):
            tbl = doc.add_table(rows=2, cols=2)
            tbl.cell(0, 0).text = "Header1"
            tbl.cell(0, 1).text = "Header2"
            tbl.cell(1, 0).text = "Val1"
            tbl.cell(1, 1).text = "Val2"
            
        raw = create_mock_docx_with_builder(build)
        res = extract_text_from_docx(raw)
        self.assertEqual(res["status"], "ok")
        self.assertIn("Header1\tHeader2", res["text"])
        self.assertIn("Val1\tVal2", res["text"])
        self.assertEqual(res["metadata"]["table_count"], 1)

    def test_docx_003_mixed_order(self):
        def build(doc):
            doc.add_paragraph("Paragraph First")
            tbl = doc.add_table(rows=1, cols=1)
            tbl.cell(0, 0).text = "Table Middle"
            doc.add_paragraph("Paragraph Last")
            
        raw = create_mock_docx_with_builder(build)
        res = extract_text_from_docx(raw)
        self.assertEqual(res["status"], "ok")
        p1 = res["text"].find("Paragraph First")
        t1 = res["text"].find("Table Middle")
        p2 = res["text"].find("Paragraph Last")
        self.assertTrue(p1 < t1 < p2, "Document order paragraph -> table -> paragraph must be preserved")

    def test_docx_004_empty_document(self):
        def build(doc):
            pass
            
        raw = create_mock_docx_with_builder(build)
        res = extract_text_from_docx(raw)
        self.assertEqual(res["status"], "ok")
        self.assertEqual(res["text"], "")
        self.assertTrue(len(res["warnings"]) > 0)

    def test_docx_005_corrupt_file(self):
        corrupt_bytes = b"PK\x03\x04this is corrupt garbage not a real zip"
        with self.assertRaises(DocxValidationError) as ctx:
            validate_docx_preflight(corrupt_bytes)
        self.assertEqual(ctx.exception.error_code, "DOCX_INVALID")

    def test_docx_006_macro_rejection(self):
        raw = create_corrupt_or_custom_zip(extra_files={"word/vbaProject.bin": b"VBA binary payload"})
        with self.assertRaises(DocxValidationError) as ctx:
            validate_docx_preflight(raw)
        self.assertEqual(ctx.exception.error_code, "UNSUPPORTED_MACRO_ENABLED_OFFICE")

    def test_docx_007_path_traversal(self):
        raw = create_corrupt_or_custom_zip(extra_files={"../evil.txt": b"traversal"})
        with self.assertRaises(DocxValidationError) as ctx:
            validate_docx_preflight(raw)
        self.assertEqual(ctx.exception.error_code, "PATH_TRAVERSAL_DETECTED")

    def test_docx_008_zip_guardrail(self):
        raw = create_corrupt_or_custom_zip()
        # Test max entries limit
        with self.assertRaises(DocxValidationError) as ctx:
            validate_docx_preflight(raw, max_entries=2)
        self.assertEqual(ctx.exception.error_code, "ZIP_GUARDRAIL_EXCEEDED")

    def test_docx_009_mime_mismatch_or_non_zip(self):
        non_zip = b"NOT_A_ZIP_HEADER_DATA"
        with self.assertRaises(DocxValidationError) as ctx:
            validate_docx_preflight(non_zip)
        self.assertEqual(ctx.exception.error_code, "DOCX_INVALID")

    def test_docx_010_determinism(self):
        def build(doc):
            doc.add_paragraph("Deterministic Text 123")
            
        raw = create_mock_docx_with_builder(build)
        res1 = extract_text_from_docx(raw)
        res2 = extract_text_from_docx(raw)
        self.assertEqual(res1["text"], res2["text"])
        self.assertEqual(res1["metadata"], res2["metadata"])


if __name__ == "__main__":
    unittest.main()
